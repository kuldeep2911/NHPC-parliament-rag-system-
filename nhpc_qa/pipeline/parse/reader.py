"""
File reader / router.

Reads the selected answer file and produces an IR `Document` (ordered blocks +
cleaned tables + language tags), choosing the parser by file type and by whether
a PDF actually has a text layer.

Routing:
    digital PDF   -> Docling (primary) ; pdfplumber fallback
    scanned PDF   -> Docling w/ OCR (rapidocr) ; flagged ocr_used
    DOCX          -> Docling ; python-docx fallback
    DOC / RTF     -> libreoffice -> DOCX, then as DOCX (flagged format_converted)
                     ; if libreoffice absent, flagged + best-effort text
    XLSX          -> openpyxl/pandas tables directly (no OCR)
    TXT           -> read directly

Every adapter is optional: if a library isn't importable, the router degrades to
the next best available reader and records a flag. It never raises to the caller;
on hard failure it returns a Document with an 'error' flag and empty content.
"""

from __future__ import annotations

import os
import re
import shutil
import subprocess
import tempfile

from .ir import Block, Document, RawTable, detect_language
from .routing import classify_pages, summarize_routing
from nhpc_qa.core.providers.models import BackendError, NotWiredError

# --- lazy capability detection ---------------------------------------------

def _has(mod: str) -> bool:
    try:
        __import__(mod)
        return True
    except Exception:
        return False


HAS_DOCLING = _has("docling")
HAS_PDFPLUMBER = _has("pdfplumber")
HAS_DOCX = _has("docx")
HAS_OPENPYXL = _has("openpyxl")
HAS_PDFIUM = _has("pypdfium2")


def render_page_png(path: str, page_no: int, scale: float = 2.0):
    """Render 1-based PDF page to PNG bytes via pypdfium2. None if unavailable."""
    if not HAS_PDFIUM:
        return None
    try:
        import io
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(path)
        try:
            page = pdf[page_no - 1]
            bitmap = page.render(scale=scale)
            pil = bitmap.to_pil()
            buf = io.BytesIO()
            pil.save(buf, format="PNG")
            return buf.getvalue()
        finally:
            pdf.close()
    except Exception:
        return None

_DOCLING_CONVERTER = None  # built once, reused (model load is expensive)


def _docling_converter(cfg=None):
    """
    Build (once) a Docling converter with IBM TableFormer explicitly configured for
    table-structure extraction. TableFormer runs on CPU and recovers rows/cols/
    spans; ACCURATE mode is used by default for the NHPC merged-header tables.
    """
    global _DOCLING_CONVERTER
    if _DOCLING_CONVERTER is None:
        from docling.document_converter import DocumentConverter, PdfFormatOption
        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import (
            PdfPipelineOptions, TableFormerMode)

        mode_name = (getattr(cfg, "tableformer_mode", "accurate") or "accurate").upper()
        mode = getattr(TableFormerMode, mode_name, TableFormerMode.ACCURATE)

        opts = PdfPipelineOptions()
        opts.do_table_structure = True
        opts.table_structure_options.mode = mode
        opts.table_structure_options.do_cell_matching = (
            getattr(cfg, "tableformer_cell_matching", True))
        try:
            _DOCLING_CONVERTER = DocumentConverter(
                format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=opts)})
        except Exception:
            # if the format-option wiring changes across Docling versions, fall back
            # to defaults (which are also TableFormer ACCURATE + cell matching).
            _DOCLING_CONVERTER = DocumentConverter()
    return _DOCLING_CONVERTER


# --- scanned detection ------------------------------------------------------

def pdf_is_scanned(path: str, threshold_per_page: int) -> bool:
    """Digital vs scanned by extractable text length per page (not by extension)."""
    if not HAS_PDFPLUMBER:
        return False
    try:
        import pdfplumber
        import warnings
        warnings.filterwarnings("ignore")
        with pdfplumber.open(path) as pdf:
            npg = len(pdf.pages)
            chars = 0
            for pg in pdf.pages[: min(npg, 8)]:
                chars += len((pg.extract_text() or "").strip())
            return chars < threshold_per_page * min(npg, 8)
    except Exception:
        return False


# --- parser adapters --------------------------------------------------------

def _docling_to_document(path: str, ocr: bool, cfg=None) -> Document:
    conv = _docling_converter(cfg)
    res = conv.convert(path)
    d = res.document
    doc = Document(parser_used="docling", ocr_used=ocr)
    # blocks in reading order via markdown export split, keeping tables separate
    try:
        page_count = len({getattr(it.prov[0], "page_no", 1)
                          for it, _ in d.iterate_items() if getattr(it, "prov", None)})
    except Exception:
        page_count = 0
    doc.page_count = page_count or 1

    # text blocks
    for text, page in _docling_text_blocks(d):
        if text.strip():
            doc.blocks.append(Block(text=text, page=page,
                                    language=detect_language(text)))
    # tables
    for ti, tb in enumerate(getattr(d, "tables", []) or [], start=1):
        grid = _docling_table_grid(tb)
        if grid:
            page = _docling_item_page(tb)
            doc.tables.append(RawTable(grid=grid, page=page, source="docling"))
    return doc


def _docling_text_blocks(d):
    """Yield (text, page) for text items, skipping table items."""
    out = []
    try:
        for item, _level in d.iterate_items():
            cls = type(item).__name__
            if "Table" in cls:
                continue
            txt = getattr(item, "text", None)
            if not txt:
                continue
            page = 1
            prov = getattr(item, "prov", None)
            if prov:
                page = getattr(prov[0], "page_no", 1) or 1
            out.append((txt, page))
    except Exception:
        # fall back to a single markdown blob
        try:
            md = d.export_to_markdown()
            for line in md.splitlines():
                if line.strip() and not line.strip().startswith("|"):
                    out.append((line, 1))
        except Exception:
            pass
    return out


def _docling_item_page(item):
    prov = getattr(item, "prov", None)
    if prov:
        return getattr(prov[0], "page_no", 1) or 1
    return 1


def _docling_table_grid(tb):
    """
    Extract a rectangular grid (list[list[str]]) from a Docling/TableFormer table.

    Uses TableFormer's predicted CELL structure (start/end row+col offsets), which
    preserves merged/spanning cells (a spanned cell's text is repeated into every
    position it covers, so multi-row/-column headers survive). Because TableFormer
    occasionally under-predicts columns on some tables, we ALSO build the flat
    dataframe grid and keep whichever preserves more content (more non-empty cells /
    columns) — never silently dropping columns.
    """
    cell_grid = _docling_cells_to_grid(tb)
    df_grid = _docling_dataframe_grid(tb)
    if cell_grid and df_grid:
        return cell_grid if _grid_richness(cell_grid) >= _grid_richness(df_grid) else df_grid
    return cell_grid or df_grid


def _grid_richness(grid):
    """
    Score a grid so we keep the better extraction. A grid whose header row is just
    positional integers (0,1,2 — a sign the header wasn't captured) is penalised so
    a grid with real header names wins even if it has slightly fewer filled cells.
    Score = (has_real_header, #columns, #non-empty cells).
    """
    if not grid:
        return (0, 0, 0)
    ncols = max((len(r) for r in grid), default=0)
    filled = sum(1 for row in grid for c in row if str(c).strip())
    header = grid[0] if grid else []
    positional = all(str(h).strip().isdigit() for h in header if str(h).strip())
    has_real_header = 0 if (positional and header) else 1
    return (has_real_header, ncols, filled)


def _docling_dataframe_grid(tb):
    try:
        df = tb.export_to_dataframe()
        header = [str(c) for c in df.columns]
        rows = [[("" if v is None else str(v)) for v in row]
                for row in df.itertuples(index=False, name=None)]
        return [header] + rows
    except Exception:
        return None


def _docling_cells_to_grid(tb):
    """Build a grid from TableFormer table_cells, repeating spanned cell values."""
    try:
        data = tb.data
        cells = data.table_cells
        if not cells:
            return None
        nrows = getattr(data, "num_rows", 0) or (
            max(c.end_row_offset_idx for c in cells))
        ncols = getattr(data, "num_cols", 0) or (
            max(c.end_col_offset_idx for c in cells))
        grid = [["" for _ in range(ncols)] for _ in range(nrows)]
        for cell in cells:
            text = (cell.text or "").strip()
            r0 = cell.start_row_offset_idx
            r1 = getattr(cell, "end_row_offset_idx", r0 + 1)
            c0 = cell.start_col_offset_idx
            c1 = getattr(cell, "end_col_offset_idx", c0 + 1)
            for r in range(r0, min(r1, nrows)):
                for c in range(c0, min(c1, ncols)):
                    # repeat spanned value into each covered position (flag handled
                    # downstream in tables.clean_table as merged_cells_present)
                    if not grid[r][c]:
                        grid[r][c] = text
        # drop trailing fully-empty rows
        while grid and not any(grid[-1]):
            grid.pop()
        return grid or None
    except Exception:
        return None


def _pdfplumber_to_document(path: str) -> Document:
    import pdfplumber
    import warnings
    warnings.filterwarnings("ignore")
    doc = Document(parser_used="pdfplumber")
    with pdfplumber.open(path) as pdf:
        doc.page_count = len(pdf.pages)
        for pi, pg in enumerate(pdf.pages, start=1):
            txt = pg.extract_text() or ""
            if txt.strip():
                doc.blocks.append(Block(text=txt, page=pi,
                                        language=detect_language(txt)))
            try:
                for t in pg.extract_tables():
                    if t and any(any(c for c in row) for row in t):
                        doc.tables.append(RawTable(grid=t, page=pi,
                                                   source="pdfplumber",
                                                   extraction_confidence="low"))
            except Exception:
                pass
    doc.add_flag("pdfplumber_fallback")
    return doc


def _docx_to_document(path: str) -> Document:
    import docx
    doc = Document(parser_used="python-docx")
    d = docx.Document(path)
    for para in d.paragraphs:
        if para.text.strip():
            doc.blocks.append(Block(text=para.text, page=1,
                                    language=detect_language(para.text)))
    for t in d.tables:
        grid = [[cell.text for cell in row.cells] for row in t.rows]
        if grid:
            doc.tables.append(RawTable(grid=grid, page=1, source="docx"))
    doc.page_count = 1
    return doc


def _xlsx_to_document(path: str) -> Document:
    import openpyxl
    doc = Document(parser_used="openpyxl")
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    for ws in wb.worksheets:
        grid = []
        for row in ws.iter_rows(values_only=True):
            grid.append([("" if v is None else str(v)) for v in row])
        grid = [r for r in grid if any(c.strip() for c in r)]
        if grid:
            doc.tables.append(RawTable(grid=grid, page=1, source="xlsx",
                                       caption=ws.title))
    doc.page_count = 1
    return doc


def _txt_to_document(path: str) -> Document:
    doc = Document(parser_used="text")
    with open(path, encoding="utf-8", errors="replace") as fh:
        txt = fh.read()
    for para in re.split(r"\n\s*\n", txt):
        if para.strip():
            doc.blocks.append(Block(text=para.strip(), page=1,
                                    language=detect_language(para)))
    doc.page_count = 1
    return doc


def _find_libreoffice(cfg):
    """Locate soffice: explicit config, then PATH, then common install dirs."""
    if cfg.libreoffice_bin and os.path.exists(cfg.libreoffice_bin):
        return cfg.libreoffice_bin
    found = shutil.which("libreoffice") or shutil.which("soffice")
    if found:
        return found
    for cand in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice", "/usr/bin/libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ):
        if os.path.exists(cand):
            return cand
    return None


def _convert_legacy(path: str, cfg) -> str | None:
    """Convert DOC/RTF to DOCX via libreoffice/soffice. Returns new path or None."""
    binary = _find_libreoffice(cfg)
    if not binary:
        return None
    outdir = tempfile.mkdtemp(prefix="nhpc_conv_")
    try:
        subprocess.run(
            [binary, "--headless", "--convert-to", "docx", "--outdir", outdir, path],
            check=True, capture_output=True, timeout=180,
        )
        stem = os.path.splitext(os.path.basename(path))[0]
        cand = os.path.join(outdir, stem + ".docx")
        return cand if os.path.exists(cand) else None
    except Exception:
        return None


# --- per-page PDF routing ---------------------------------------------------

def _read_pdf_per_page(path: str, cfg, provider, tracer) -> Document:
    """
    Classify each page (digital/scanned/image_based) and route it:
        digital     -> parse text+tables (Docling primary, pdfplumber fallback)
        scanned     -> provider.ocr_image(rendered page)
        image_based -> provider.parse_visual(rendered page)
    Pages are reassembled in reading order. Every decision is recorded on
    doc.page_routing and (if a tracer is given) emitted as a routing trace step.
    """
    routes = classify_pages(path, cfg)
    summary = summarize_routing(routes)

    # Parse all digital content ONCE with the primary parser, indexed by page.
    digital_pages = {r.page for r in routes if r.page_type == "digital"}
    base = _parse_digital_pdf(path, cfg)  # a Document with page-tagged blocks/tables

    doc = Document(parser_used=base.parser_used)
    doc.page_count = summary["n_pages"]
    doc.models_used = {}

    for r in routes:
        if tracer:
            tracer.step("routing", {"page": r.page, "decision": r.to_dict()})
        use_nim_tables = _is_nim_provider(provider)
        if r.page_type == "digital":
            for b in base.blocks:
                if b.page == r.page:
                    doc.blocks.append(b)
            if not use_nim_tables:  # NIM path supplies tables separately below
                for t in base.tables:
                    if t.page == r.page:
                        doc.tables.append(t)
        elif r.page_type == "scanned":
            _route_ocr(doc, path, r, cfg, provider)
        elif r.page_type == "image_based":
            _route_visual(doc, path, r, cfg, provider)

        # When the NVIDIA NeMo Retriever NIMs are the backend, extract tables on
        # this page with the page-elements + table-structure + OCR NIM pipeline
        # (replaces Docling/TableFormer for tables). Digital text still comes from
        # the digital parse above; only the TABLES are taken from the NIMs.
        if _is_nim_provider(provider):
            _route_nim_tables(doc, path, r, cfg, provider)

    # if digital parsing found content on pages we didn't mark digital (rare), keep it
    if not doc.blocks and not doc.tables and (base.blocks or base.tables):
        doc.blocks, doc.tables = base.blocks, base.tables

    doc.page_routing = [r.to_dict() for r in routes]
    if summary["mixed"]:
        doc.add_flag("mixed_page_types")
    if summary["scanned"]:
        doc.add_flag("ocr_used")
        doc.ocr_used = True
    if summary["image_based"]:
        doc.add_flag("visual_used")
        doc.visual_used = True
    for fl in base.flags:
        doc.add_flag(fl)
    return doc


def _parse_digital_pdf(path: str, cfg) -> Document:
    """Text+table parse of the (digital portions of the) PDF, page-tagged."""
    if cfg.prefer_docling and HAS_DOCLING:
        try:
            d = _docling_to_document(path, ocr=False, cfg=cfg)
            if d.blocks or d.tables:
                return d
        except Exception:
            pass
    if HAS_PDFPLUMBER:
        return _pdfplumber_to_document(path)
    d = Document(parser_used="none")
    d.add_flag("no_pdf_parser")
    return d


def _route_ocr(doc, path, r, cfg, provider):
    """OCR a scanned page image through the provider; append as a text block."""
    if provider is None or not cfg.enable_ocr:
        doc.add_flag("ocr_unavailable")
        return
    img = render_page_png(path, r.page)
    if img is None:
        doc.add_flag("page_render_unavailable")
        return
    try:
        text = provider.ocr_image(img, lang=cfg.ocr_lang)
        if text and text.strip():
            doc.blocks.append(Block(text=text.strip(), page=r.page,
                                    language=detect_language(text)))
            doc.models_used["ocr"] = provider.model_for("ocr")
    except NotWiredError:
        doc.add_flag("ocr_backend_not_wired")
    except (BackendError, Exception):
        doc.add_flag("ocr_failed")


def _route_visual(doc, path, r, cfg, provider):
    """Send an image-based page to the provider VLM; append extracted text."""
    if provider is None:
        doc.add_flag("visual_unavailable")
        return
    img = render_page_png(path, r.page)
    if img is None:
        doc.add_flag("page_render_unavailable")
        return
    prompt = ("Transcribe all text visible in this page image verbatim in reading "
              "order, preserving Hindi/Devanagari exactly. If a table is present, "
              "render it as rows. Output only the transcribed content.")
    try:
        res = provider.parse_visual(img, prompt)
        text = res.get("text") if isinstance(res, dict) else str(res)
        if text and text.strip():
            doc.blocks.append(Block(text=text.strip(), page=r.page,
                                    language=detect_language(text)))
            doc.models_used["visual"] = provider.model_for("visual")
    except NotWiredError:
        doc.add_flag("visual_backend_not_wired")
    except (BackendError, Exception):
        doc.add_flag("visual_failed")


# --- NVIDIA NeMo Retriever NIM table pipeline -------------------------------

def _is_nim_provider(provider) -> bool:
    return getattr(provider, "kind", None) == "nvidia" and hasattr(
        provider, "detect_table_structure")


def _route_nim_tables(doc, path, r, cfg, provider):
    """
    Extract tables on this page using the NeMo Retriever NIMs:
      1. page-elements NIM  -> locate 'table' regions on the page
      2. for each table region, table-structure NIM -> cell/row/column boxes
      3. OCR NIM on the table crop -> text per cell, assembled into a grid
    Each recovered table is appended to doc.tables as a RawTable(source='nim').
    """
    img = render_page_png(path, r.page)
    if img is None:
        doc.add_flag("page_render_unavailable")
        return
    try:
        elements = provider.detect_page_elements(img)
    except NotWiredError:
        doc.add_flag("nim_not_wired")
        return
    except (BackendError, Exception):
        doc.add_flag("nim_page_elements_failed")
        return

    table_boxes = (elements or {}).get("table", [])
    if not table_boxes:
        return
    doc.models_used["page_elements"] = provider.model_for("page_elements")

    from PIL import Image
    import io
    try:
        page_img = Image.open(io.BytesIO(img)).convert("RGB")
    except Exception:
        doc.add_flag("nim_crop_failed")
        return
    W, H = page_img.size

    for ti, box in enumerate(table_boxes, start=1):
        crop = _crop_norm(page_img, box, W, H)
        if crop is None:
            continue
        crop_bytes = _png_bytes(crop)
        try:
            struct = provider.detect_table_structure(crop_bytes)
        except (NotWiredError, BackendError, Exception):
            doc.add_flag("nim_table_structure_failed")
            continue
        grid = _nim_build_grid(provider, crop, struct)
        if grid:
            doc.tables.append(RawTable(grid=grid, page=r.page, source="nim"))
            doc.models_used["table_structure"] = provider.model_for("table_structure")
            doc.add_flag("nim_table_extracted")


def _crop_norm(img, box, W, H):
    x0 = int(max(0.0, box.get("x_min", 0.0)) * W)
    y0 = int(max(0.0, box.get("y_min", 0.0)) * H)
    x1 = int(min(1.0, box.get("x_max", 1.0)) * W)
    y1 = int(min(1.0, box.get("y_max", 1.0)) * H)
    if x1 - x0 < 4 or y1 - y0 < 4:
        return None
    return img.crop((x0, y0, x1, y1))


def _png_bytes(pil_img):
    import io
    buf = io.BytesIO()
    pil_img.save(buf, format="PNG")
    return buf.getvalue()


def _nim_build_grid(provider, table_img, struct):
    """
    Build a cell grid from table-structure row/column boxes + per-cell OCR.
    Rows and columns from the table-structure NIM define a lattice; each lattice
    cell is OCR'd (via the OCR NIM) to get its text.
    """
    rows = sorted(struct.get("row", []), key=lambda b: b.get("y_min", 0.0))
    cols = sorted(struct.get("column", []), key=lambda b: b.get("x_min", 0.0))
    if not rows or not cols:
        # fall back to explicit 'cell' boxes if row/column not provided
        cells = struct.get("cell", [])
        if not cells:
            return None
        return _grid_from_cells(provider, table_img, cells)

    W, H = table_img.size
    grid = []
    for rb in rows:
        row_vals = []
        for cb in cols:
            x0 = int(cb.get("x_min", 0.0) * W); x1 = int(cb.get("x_max", 1.0) * W)
            y0 = int(rb.get("y_min", 0.0) * H); y1 = int(rb.get("y_max", 1.0) * H)
            if x1 - x0 < 2 or y1 - y0 < 2:
                row_vals.append("")
                continue
            cell_crop = table_img.crop((x0, y0, x1, y1))
            try:
                txt = provider.ocr_image(_png_bytes(cell_crop))
            except Exception:
                txt = ""
            row_vals.append(re.sub(r"\s+", " ", txt).strip())
        grid.append(row_vals)
    return grid or None


def _grid_from_cells(provider, table_img, cells):
    """Assemble a grid from individual cell boxes by clustering their positions."""
    W, H = table_img.size
    # cluster distinct row (y) and column (x) coordinates
    ys = sorted({round(c.get("y_min", 0.0), 2) for c in cells})
    xs = sorted({round(c.get("x_min", 0.0), 2) for c in cells})
    row_of = {y: i for i, y in enumerate(ys)}
    col_of = {x: i for i, x in enumerate(xs)}
    grid = [["" for _ in xs] for _ in ys]
    for c in cells:
        ri = row_of.get(round(c.get("y_min", 0.0), 2))
        ci = col_of.get(round(c.get("x_min", 0.0), 2))
        if ri is None or ci is None:
            continue
        x0 = int(c.get("x_min", 0.0) * W); x1 = int(c.get("x_max", 1.0) * W)
        y0 = int(c.get("y_min", 0.0) * H); y1 = int(c.get("y_max", 1.0) * H)
        if x1 - x0 < 2 or y1 - y0 < 2:
            continue
        try:
            txt = provider.ocr_image(_png_bytes(table_img.crop((x0, y0, x1, y1))))
        except Exception:
            txt = ""
        grid[ri][ci] = re.sub(r"\s+", " ", txt).strip()
    return grid or None


# --- main entry -------------------------------------------------------------

def read_document(path: str, cfg, provider=None, tracer=None) -> Document:
    """
    Route `path` to the right parser and return an IR Document. Never raises.

    PDFs are routed PER PAGE (see routing.classify_pages): digital pages are parsed
    for text/tables, scanned pages are OCR'd, image-based pages go through the VLM.
    `provider` supplies ocr_image()/parse_visual(); if None, those pages are flagged
    and skipped rather than crashing. Per-page decisions land on doc.page_routing.
    """
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _read_pdf_per_page(path, cfg, provider, tracer)

        if ext == ".docx":
            if cfg.prefer_docling and HAS_DOCLING:
                try:
                    return _docling_to_document(path, ocr=False, cfg=cfg)
                except Exception:
                    pass
            return _docx_to_document(path)

        if ext in (".doc", ".rtf"):
            newp = _convert_legacy(path, cfg)
            if newp:
                doc = (_docling_to_document(newp, ocr=False, cfg=cfg)
                       if cfg.prefer_docling and HAS_DOCLING else _docx_to_document(newp))
                doc.format_converted = True
                doc.add_flag("format_converted")
                return doc
            # no converter: best-effort raw text, flagged
            doc = Document(parser_used="raw")
            doc.format_converted = False
            doc.add_flag("format_conversion_unavailable")
            try:
                with open(path, "rb") as fh:
                    raw = fh.read()
                text = raw.decode("utf-8", errors="ignore")
                text = re.sub(r"[^\x09\x0a\x0d\x20-\x7eऀ-ॿ]", " ", text)
                text = re.sub(r"\s+", " ", text)
                if len(text.strip()) > 40:
                    doc.blocks.append(Block(text=text.strip()[:20000], page=1,
                                            language=detect_language(text)))
            except Exception:
                pass
            return doc

        if ext == ".xlsx":
            return _xlsx_to_document(path)

        if ext == ".txt":
            return _txt_to_document(path)

        # unknown extension
        doc = Document(parser_used="unknown")
        doc.add_flag("unsupported_filetype")
        return doc

    except Exception as e:  # never propagate — one bad doc must not abort the run
        doc = Document(parser_used="error")
        doc.add_flag("read_error")
        doc.add_flag(f"error:{type(e).__name__}")
        return doc
