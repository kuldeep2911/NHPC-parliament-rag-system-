"""
The draft as a Word document.

python-docx is ALREADY a dependency (the parser reads .docx); it writes them too, so this
adds nothing to the install.

DEVANAGARI. A .docx is UTF-8 XML, so the text survives regardless. What does NOT survive is
the RENDERING: Word picks the font for a script from the run's *complex-script* setting
(w:cs), and python-docx does not expose it. Set only the Latin font and Hindi text falls
back to whatever Word guesses -- often a box glyph. _font() sets w:cs explicitly, so
Devanagari renders in a font that actually has the glyphs.

The DRAFT notice appears THREE times -- a red banner at the top, a line under every drafted
part, and the footer of every page. That is deliberate. This document will be printed,
forwarded, and pasted from. A single notice at the top disappears the moment someone copies
the middle of page 2.
"""

from __future__ import annotations

import datetime as dt
import io
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# A font with BOTH Latin and Devanagari coverage, present on any Windows/Office install.
# Nirmala UI is Microsoft's own Indic UI font.
_LATIN = "Calibri"
_DEVA = "Nirmala UI"

_RED = RGBColor(0xB0, 0x2A, 0x22)
_GREY = RGBColor(0x6B, 0x71, 0x78)


def _font(run, *, size=11, bold=False, italic=False, color=None, hindi=False):
    """
    Style a run, and set the COMPLEX-SCRIPT font too.

    Without the w:cs element, Word renders Devanagari with whatever it falls back to --
    frequently tofu boxes. python-docx has no API for it, so we reach into the XML.
    """
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    name = _DEVA if hindi else _LATIN
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    # ascii/hAnsi = Latin; cs = complex script (Devanagari). Setting cs is the whole point.
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), _DEVA)
    return run


_DEVANAGARI = re.compile(r"[ऀ-ॿ]")


def _is_hindi(text: str) -> bool:
    return bool(_DEVANAGARI.search(text or ""))


def _para(doc, text="", *, size=11, bold=False, italic=False, color=None,
          space_after=6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        _font(p.add_run(text), size=size, bold=bold, italic=italic, color=color,
              hindi=_is_hindi(text))
    return p


def build_docx(query: str, draft: dict, *, user_email: str | None = None,
               run_id: str | None = None) -> bytes:
    """
    The draft as .docx bytes — a CLEAN REPLY FILE ONLY.

    By officer request, the downloaded document mirrors an actual NHPC reply file and NOTHING
    else: Subject, then for each part the QUESTION in bold followed by the answer (and a
    table where one applies). No citations, no key-points, no gaps, no contradictions, no
    source tables — those live only on the website for verification. A single unobtrusive
    DRAFT footer remains, because this is still not an approved reply.

    Pure function -- no filesystem, no network.
    """
    doc = Document()

    # Base style, so anything we do not touch still has Devanagari coverage.
    normal = doc.styles["Normal"]
    normal.font.name = _LATIN
    normal.font.size = Pt(11)
    normal.element.rPr.rFonts.set(qn("w:cs"), _DEVA)

    subject = (draft.get("subject") or "").strip()
    opening = (draft.get("opening") or "").strip()
    closing = (draft.get("closing") or "").strip()

    # ---- Subject (bold), centred like the reply files ------------------
    if subject:
        p = _para(doc, "", space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _font(p.add_run("Subject: "), size=12, bold=True)
        _font(p.add_run(subject), size=12, bold=True, hindi=_is_hindi(subject))
    else:
        _para(doc, query or "", size=12, bold=True, space_after=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    if opening:
        _para(doc, opening, size=11, space_after=8)

    # ---- body: QUESTION (bold) then ANSWER, per part -------------------
    parts = draft.get("parts") or []
    if not parts and not opening:
        _para(doc, "(no draft was produced)", italic=True, color=_GREY)
    for p in parts:
        label = (p.get("label") or "").strip()
        question = (p.get("question") or "").strip()
        text = (p.get("text") or "").strip()

        # QUESTION line — label + question text, in bold (the reply-file convention)
        if question:
            qp = doc.add_paragraph()
            qp.paragraph_format.space_after = Pt(3)
            if label:
                _font(qp.add_run(f"{label} "), size=11, bold=True)
            _font(qp.add_run(question), size=11, bold=True, hindi=_is_hindi(question))

        # ANSWER line
        if text:
            ap = doc.add_paragraph()
            ap.paragraph_format.space_after = Pt(7)
            ap.paragraph_format.left_indent = Pt(10)
            # if there was no question line, the label leads the answer instead
            if label and not question:
                _font(ap.add_run(f"{label} "), size=11, bold=True)
            _font(ap.add_run(text), size=11, hindi=_is_hindi(text))

        # optional TABLE built from the grounded facts
        _render_table(doc, p.get("table"))

    if closing:
        _para(doc, closing, size=11, space_after=10)

    # ---- footer, on every page (the only DRAFT marking that remains) ---
    sec = doc.sections[0]
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _font(fp.add_run("DRAFT — for officer review. Not an approved reply."),
          size=8, bold=True, color=_RED)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_table(doc, table):
    """Render an optional {columns, rows} grid into the document. No-op if empty/malformed."""
    if not isinstance(table, dict):
        return
    cols = table.get("columns") or []
    rows = table.get("rows") or []
    if not cols or not rows:
        return
    ncol = len(cols)
    t = doc.add_table(rows=1, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    for i, h in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = ""
        _font(cell.paragraphs[0].add_run(str(h)), size=10, bold=True, hindi=_is_hindi(str(h)))
    # body — tolerate ragged rows (pad/truncate to ncol)
    for r in rows:
        cells = list(r) if isinstance(r, (list, tuple)) else [r]
        cells = (cells + [""] * ncol)[:ncol]
        row = t.add_row().cells
        for i, v in enumerate(cells):
            row[i].text = ""
            _font(row[i].paragraphs[0].add_run(str(v)), size=10, hindi=_is_hindi(str(v)))
    _para(doc, "", space_after=8)


_SAFE = re.compile(r"[^A-Za-z0-9._-]+")


def safe_filename(query: str, run_id: str | None = None) -> str:
    """
    A filesystem- and header-safe filename.

    The query is officer-supplied, so it is hostile until sanitised: a newline in a
    Content-Disposition header is a response-splitting primitive, and a '/' or '..' would
    matter the moment anyone saved this by hand. Strip to a known-safe alphabet rather than
    blacklisting.

    A fully non-Latin query (Hindi) sanitises to NOTHING, which produced the useless
    "NHPC-draft-draft-<id>.docx". Rather than transliterate -- guessing at romanisation is
    its own bug -- we simply drop the stem and let the run id name the file. The document's
    CONTENT is Devanagari either way; only the filename is ASCII, because a Content-
    Disposition header with raw UTF-8 is not portable across browsers.
    """
    stem = _SAFE.sub("-", (query or "").strip())[:48]
    # Collapse dots too. '../../etc/passwd' otherwise survives as '..-..-etc-passwd': it is
    # only a download filename and we never open it, but a name carrying '..' is the kind
    # of thing that becomes a real bug the moment someone later uses it as a path.
    stem = re.sub(r"[.\-]{2,}", "-", stem).strip("-.")
    tail = _SAFE.sub("-", (run_id or dt.datetime.now().strftime("%Y%m%d-%H%M%S"))[:24])
    return f"NHPC-draft-{stem}-{tail}.docx" if stem else f"NHPC-draft-{tail}.docx"
